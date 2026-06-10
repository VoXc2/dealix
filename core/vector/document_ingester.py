from __future__ import annotations

import asyncio
import logging
import os
import tempfile
from pathlib import Path
from urllib.parse import urlparse

import aiohttp
import markdown
from bs4 import BeautifulSoup

from .client import VectorClient
from .embedding_pipeline import EmbeddingPipeline
from .schemas import Document, DocumentType, IngestResult, VectorPoint

logger = logging.getLogger(__name__)


class DocumentIngester:
    SUPPORTED_TYPES = [".pdf", ".docx", ".html", ".md", ".txt"]

    def __init__(
        self,
        vector_client: VectorClient,
        embedding_pipeline: EmbeddingPipeline,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ):
        self.vector_client = vector_client
        self.embedding = embedding_pipeline
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    async def ingest(self, file_path: str, collection: str) -> IngestResult:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_TYPES:
            return IngestResult(
                collection=collection,
                source=file_path,
                chunks_count=0,
                document_ids=[],
                success=False,
                error=f"Unsupported file type: {ext}. Supported: {self.SUPPORTED_TYPES}",
            )

        try:
            content = await self._read_file(file_path, ext)
            doc_type = self._ext_to_doc_type(ext)
            doc = Document(
                content=content,
                doc_type=doc_type,
                source=file_path,
                metadata={"file_path": file_path, "file_type": ext},
            )
            return await self._ingest_document(doc, collection)
        except Exception as e:
            logger.exception("Failed to ingest file: %s", file_path)
            return IngestResult(
                collection=collection,
                source=file_path,
                chunks_count=0,
                document_ids=[],
                success=False,
                error=str(e),
            )

    async def ingest_url(self, url: str, collection: str) -> IngestResult:
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=aiohttp.ClientTimeout(total=30)) as resp:
                    resp.raise_for_status()
                    html = await resp.text()

            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            content = soup.get_text(separator="\n", strip=True)

            doc = Document(
                content=content,
                doc_type=DocumentType.HTML,
                source=url,
                metadata={"url": url, "title": soup.title.string if soup.title else ""},
            )
            return await self._ingest_document(doc, collection)
        except Exception as e:
            logger.exception("Failed to ingest URL: %s", url)
            return IngestResult(
                collection=collection,
                source=url,
                chunks_count=0,
                document_ids=[],
                success=False,
                error=str(e),
            )

    async def ingest_text(
        self,
        text: str,
        collection: str,
        source: str | None = None,
    ) -> IngestResult:
        doc = Document(
            content=text,
            doc_type=DocumentType.TEXT,
            source=source or "direct_input",
            metadata={"source": source or "direct_input"},
        )
        return await self._ingest_document(doc, collection)

    async def _ingest_document(
        self,
        doc: Document,
        collection: str,
    ) -> IngestResult:
        await self.vector_client.create_collection(collection)

        doc_embeddings = await self.embedding.embed_document(doc)

        points = [
            VectorPoint(
                vector=de.embedding,
                payload={
                    "content": de.content,
                    "document_id": de.document_id,
                    "chunk_index": de.chunk_index,
                    **de.metadata,
                },
                metadata=de.metadata,
            )
            for de in doc_embeddings
        ]

        await self.vector_client.upsert(collection, points)

        return IngestResult(
            collection=collection,
            source=doc.source or "unknown",
            chunks_count=len(points),
            document_ids=[str(de.document_id) for de in doc_embeddings],
        )

    async def _read_file(self, file_path: str, ext: str) -> str:
        if ext == ".pdf":
            return await self._read_pdf(file_path)
        elif ext == ".docx":
            return await self._read_docx(file_path)
        elif ext == ".html":
            return await self._read_html(file_path)
        elif ext == ".md":
            return await self._read_markdown(file_path)
        elif ext == ".txt":
            return await self._read_text(file_path)
        raise ValueError(f"Unsupported extension: {ext}")

    async def _read_pdf(self, file_path: str) -> str:
        try:
            import pypdf
        except ImportError:
            raise ImportError("pypdf is required for PDF ingestion: pip install pypdf")
        text_parts = []
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts)

    async def _read_docx(self, file_path: str) -> str:
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for DOCX ingestion: pip install python-docx")
        doc = docx.Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)

    async def _read_html(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            html = f.read()
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)

    async def _read_markdown(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        html = markdown.markdown(content)
        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text(separator="\n", strip=True)

    async def _read_text(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _ext_to_doc_type(self, ext: str) -> DocumentType:
        mapping = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".html": DocumentType.HTML,
            ".md": DocumentType.MARKDOWN,
            ".txt": DocumentType.TEXT,
        }
        return mapping.get(ext, DocumentType.TEXT)
